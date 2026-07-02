"""Draft → {Markdown, HTML, .docx} rendering.

All formats start from the same assembled Markdown (title, opening hook,
then each section's prose) produced by SqlDraftStore.assemble_markdown, so
the three exports stay structurally consistent.
"""

from __future__ import annotations

import html as _html
import json
import re
from io import BytesIO

import yaml

from blogforge.drafts.models import Draft
from blogforge.drafts.sql_store import SqlDraftStore

EXPORT_FORMATS = ("md", "html", "docx")

# An FAQ block the GEO fix appends: a "### FAQ" heading, then **Question**/answer
# pairs. Parsed so exports can emit FAQPage schema for AI answer engines.
_FAQ_HEADING_RE = re.compile(
    r"(?im)^#{2,4}\s*(?:faqs?|frequently asked|common questions|q ?& ?a|q and a)\b.*$"
)
_FAQ_QA_RE = re.compile(r"(?ms)^\*\*(?P<q>.+?)\*\*\s*\n+(?P<a>.+?)(?=\n\s*\*\*|\Z)")


def extract_faqs(draft: Draft) -> list[tuple[str, str]]:
    """Pull (question, answer) pairs out of any FAQ block in the draft, so the
    HTML export can emit FAQPage schema. Empty when there's no FAQ."""
    pairs: list[tuple[str, str]] = []
    for section in draft.sections:
        m = _FAQ_HEADING_RE.search(section.content_md)
        if not m:
            continue
        for qa in _FAQ_QA_RE.finditer(section.content_md[m.end() :]):
            q = qa.group("q").strip()
            a = " ".join(qa.group("a").split()).strip()
            if q and a:
                pairs.append((q, a))
    return pairs


def frontmatter_block(draft: Draft) -> str:
    """A YAML frontmatter block (title / date / lastmod / pack / tags) for
    static-site generators. `date` is when the post was created (published),
    `lastmod` when it last changed — the freshness signal AI engines and SSGs
    both read. Trailing newline so body markdown follows cleanly."""
    data: dict[str, object] = {
        "title": draft.title or draft.idea.topic,
        "date": draft.created_at.date().isoformat(),
        "lastmod": draft.updated_at.date().isoformat(),
        "pack": draft.idea.pack_slug,
    }
    if draft.tags:
        data["tags"] = list(draft.tags)
    if draft.hero_image_key:
        data["image"] = draft.hero_image_key
    dumped = yaml.safe_dump(data, sort_keys=False, allow_unicode=True).strip()
    return f"---\n{dumped}\n---\n\n"


def json_ld(draft: Draft, author: str | None = None) -> str:
    """Article + (when present) FAQPage JSON-LD `<script>` blocks. Won't itself
    drive citations, but earns rich results and Bing's index (which ChatGPT
    Search partly relies on), and carries the freshness dates."""
    article: dict[str, object] = {
        "@context": "https://schema.org",
        "@type": "Article",
        "headline": draft.title or draft.idea.topic or "Untitled",
        "datePublished": draft.created_at.date().isoformat(),
        "dateModified": draft.updated_at.date().isoformat(),
    }
    if author:
        article["author"] = {"@type": "Person", "name": author}
    blocks: list[dict[str, object]] = [article]
    faqs = extract_faqs(draft)
    if faqs:
        blocks.append(
            {
                "@context": "https://schema.org",
                "@type": "FAQPage",
                "mainEntity": [
                    {
                        "@type": "Question",
                        "name": q,
                        "acceptedAnswer": {"@type": "Answer", "text": a},
                    }
                    for q, a in faqs
                ],
            }
        )
    return "\n".join(
        f'<script type="application/ld+json">{json.dumps(b, ensure_ascii=False)}</script>'
        for b in blocks
    )


def to_markdown(draft: Draft, *, frontmatter: bool = False) -> str:
    body = SqlDraftStore.assemble_markdown(draft)
    return frontmatter_block(draft) + body if frontmatter else body


def to_html(draft: Draft, *, hero_data_uri: str | None = None, author: str | None = None) -> str:
    """A standalone, self-styled HTML document. When `hero_data_uri` is given
    (a base64 data: URI of the hero image), it's embedded at the top so the
    exported file stays self-contained. Embeds Article/FAQPage JSON-LD and a
    visible "Updated {month}" line — the GEO freshness + citability signals."""
    import markdown as md_lib  # type: ignore[import-untyped]

    body_md = SqlDraftStore.assemble_markdown(draft)
    body_html = md_lib.markdown(body_md, extensions=["extra", "sane_lists", "smarty"])
    title = _html.escape(draft.title or draft.idea.topic or "Untitled")
    hero = (
        f'<figure class="hero"><img src="{hero_data_uri}" alt="{title}"></figure>\n'
        if hero_data_uri
        else ""
    )
    byline = f'<p class="byline">Updated {draft.updated_at.strftime("%B %Y")}</p>\n'
    return _HTML_TEMPLATE.format(
        title=title, hero=hero, byline=byline, body=body_html, jsonld=json_ld(draft, author)
    )


def to_docx(draft: Draft) -> bytes:
    """A Word document. Block-level structure (headings, paragraphs, bullets)
    is preserved; inline emphasis is flattened to plain text."""
    from docx import Document

    doc = Document()
    if draft.title or draft.idea.topic:
        doc.add_heading(draft.title or draft.idea.topic, level=0)
    if draft.outline and draft.outline.opening_hook.strip():
        doc.add_paragraph(_strip_inline(draft.outline.opening_hook.strip()))
    for section in draft.sections:
        doc.add_heading(section.title, level=1)
        for block in _paragraphs(section.content_md):
            if block.startswith(("- ", "* ")):
                doc.add_paragraph(_strip_inline(block[2:].strip()), style="List Bullet")
            else:
                doc.add_paragraph(_strip_inline(block))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── helpers ──────────────────────────────────────────────────────────


def _paragraphs(md: str) -> list[str]:
    """Split section markdown into blank-line-separated blocks, dropping
    leading section headers (the title is added as a docx heading already)."""
    blocks: list[str] = []
    for raw in re.split(r"\n\s*\n", md.strip()):
        block = raw.strip()
        if not block or block.startswith("#"):
            continue
        blocks.append(block)
    return blocks


_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_EMPHASIS_RE = re.compile(r"(\*\*|__|\*|_|`)")


def _strip_inline(text: str) -> str:
    """Flatten inline markdown (links → text, drop emphasis/code markers)."""
    text = _LINK_RE.sub(r"\1", text)
    text = _EMPHASIS_RE.sub("", text)
    return text.replace("\n", " ").strip()


_HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
{jsonld}
<style>
  body {{ max-width: 42rem; margin: 3rem auto; padding: 0 1.25rem;
    font: 18px/1.7 Georgia, "Times New Roman", serif; color: #1a1a1a; }}
  p.byline {{ font: 0.8rem/1.4 ui-sans-serif, system-ui, sans-serif;
    text-transform: uppercase; letter-spacing: 0.05em; color: #777; margin: 0 0 1.5rem; }}
  h1 {{ font-size: 2.2rem; line-height: 1.15; margin: 0 0 1.5rem; }}
  h2 {{ font-size: 1.5rem; margin: 2.5rem 0 0.75rem; }}
  p {{ margin: 0 0 1.1rem; }}
  a {{ color: #2647c0; }}
  blockquote {{ margin: 1.1rem 0; padding-left: 1rem; border-left: 3px solid #ccc; color: #555; }}
  code {{ font: 0.9em ui-monospace, Menlo, monospace; background: #f3f3f3; padding: 0.1em 0.3em; }}
  figure.hero {{ margin: 0 0 2rem; }}
  figure.hero img {{ width: 100%; height: auto; border-radius: 8px; display: block; }}
</style>
</head>
<body>
{hero}{byline}{body}
</body>
</html>
"""
